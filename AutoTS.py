import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import datetime

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Cm

#文件中读取数据
excel_file = 'SDAC_template_field.xlsx'  # Excel文件路径
#df = pd.read_excel(excel_file)  # 读取Excel文件

# 读取第一部分：用于Word占位符替换的数据
fields_df = pd.read_excel(excel_file, sheet_name='Sheet1', usecols="A:B").dropna()

# 读取第二部分：用于插入Word表格的数据
table_df = pd.read_excel(excel_file, sheet_name='Sheet1', usecols="D:E").dropna(how='all')

# Step 2: 将第一部分数据转换为字典，并格式化数值和日期
data = {}
for index, row in fields_df.iterrows():
    #print(f"Key: {index}, Value: {row}, Type: {type(row)}")
    key = row['Field Name']
    value = row['Value']

    # 如果是浮点数，确保保留小数点后4位
    if isinstance(value, float):
        data[key] = f"{value:.4f}"
    elif isinstance(value, int):
        data[key] = f"{value:,.0f}"  # 格式化为带逗号的整数
    # 如果是日期类型，保持格式为 '24-Aug-2024'
    elif isinstance(value, datetime.datetime):
        print(value)
        data[key] = value.strftime('%d-%b-%Y')
    else:
        data[key] = str(value)

print(data)

# Step 3: 读取Word模板
doc = Document('SDAC_template.docx')  # Word模板路径

# Step 4: 替换段落和表格中的占位符，同时保留格式
def replace_text_in_paragraph(paragraph, data):
    for key, value in data.items():
        full_text = ''.join(run.text for run in paragraph.runs)
        if f'{{{{ {key} }}}}' in full_text:
            new_text = full_text.replace(f'{{{{ {key} }}}}', value)
            # 清除原始段落内容
            for run in paragraph.runs:
                run.text = ''
            # 按原始格式重新插入文本
            paragraph.runs[0].text = new_text
            # 重新应用格式
            for i, run in enumerate(paragraph.runs):
                run.text = new_text[i: i + len(run.text)]

def replace_text_in_table(table, data):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph, data)

# 替换段落中的占位符，同时保留格式
for paragraph in doc.paragraphs:
    replace_text_in_paragraph(paragraph, data)

# 替换表格中的占位符，同时保留格式
for table in doc.tables:
    replace_text_in_table(table, data)


# Step 5: 添加边框的函数
def set_table_border(table):
    tbl = table._element
    tbl_pr = tbl.tblPr

    # 创建边框元素
    tbl_borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')  # 线条类型
        border.set(qn('w:sz'), '4')        # 线条大小
        border.set(qn('w:space'), '0')     # 间距
        border.set(qn('w:color'), '000000')  # 线条颜色
        tbl_borders.append(border)

    tbl_pr.append(tbl_borders)

# Step 6: 在指定位置插入表格并填充内容
def set_column_width(table, column_idx, width):
    for row in table.rows:
        row.cells[column_idx].width = width

def set_row_height(row, height):
    for cell in row.cells:
        cell.height = height

def set_table_alignment(table, alignment):
    """
    设置表格的对齐方式
    alignment 参数可以是 WD_TABLE_ALIGNMENT.LEFT, WD_TABLE_ALIGNMENT.CENTER, WD_TABLE_ALIGNMENT.RIGHT
    """
    table.alignment = alignment

def insert_table_at_paragraph(paragraph, df, alignment=WD_TABLE_ALIGNMENT.CENTER):
    table = doc.add_table(rows=1, cols=2)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = df.columns[0]
    hdr_cells[1].text = df.columns[1]

    for i, row in df.iterrows():
        row_cells = table.add_row().cells
        for j, cell_value in enumerate(row):
            if isinstance(cell_value, (int, float)):
                row_cells[j].text = f"{cell_value:,.0f}"
            elif isinstance(cell_value, pd.Timestamp):
                row_cells[j].text = cell_value.strftime('%d-%b-%Y')
            else:
                row_cells[j].text = str(cell_value)

    set_table_border(table)
    set_table_alignment(table, alignment)

    # 设置列宽为4厘米
    set_column_width(table, 0, Cm(4))
    set_column_width(table, 1, Cm(4))

    # 设置行高为1厘米
    for row in table.rows:
        set_row_height(row, Cm(1))
    paragraph._element.addnext(table._element)


target_text = "{{ insert_table_here }}"
for paragraph in doc.paragraphs:
    if target_text in paragraph.text:
        insert_table_at_paragraph(paragraph, table_df, alignment=WD_TABLE_ALIGNMENT.RIGHT)
        paragraph.text = ""
        break

# Step 6: 保存填充后的文档
doc.save('filled_termsheet.docx')