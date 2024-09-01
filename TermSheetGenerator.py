import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import datetime

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Cm


class TSGenerator():
    def __init__(self):
        self.termSheetPath = ''
        self.doc = Document(self.termSheetPath)
        self.termPath = ''
        self.fieldDF = pd.DataFrame()
        self.tableDF = pd.DataFrame()
        self.replaceDict  = self.getReplaceDict()

    def getReplaceDict(self):
        dataDict = {}
        for index, row in self.fieldDF.iterrows():
            #print(f"Key: {index}, Value: {row}, Type: {type(row)}")
            key = row['Field Name']
            value = row['Value']

            # 如果是浮点数，确保保留小数点后4位
            if isinstance(value, float):
                dataDict[key] = f"{value:.4f}"
            elif isinstance(value, int):
                dataDict[key] = f"{value:,.0f}"  # 格式化为带逗号的整数
            # 如果是日期类型，保持格式为 '24-Aug-2024'
            elif isinstance(value, datetime.datetime):
                dataDict[key] = value.strftime('%d-%b-%Y')
            else:
                dataDict[key] = str(value)
        return dataDict
    
    def replace_text_in_paragraph(self,paragraph):
        for key, value in self.replaceDict.items():
            full_text = ''.join(run.text for run in paragraph.runs)
            if f'{{{{ {key} }}}}' in full_text:
                new_text = full_text.replace(f'{{{{ {key} }}}}', value)
                # 清除原始段落内容
                for run in paragraph.runs:
                    run.text = ''
                # 按原始格式重新插入文本
                paragraph.runs[0].text = new_text
                # 重新应用格式
                for i, run in enumerate(paragraph.runs):
                    run.text = new_text[i: i + len(run.text)]

    def replace_text_in_table(self,table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self.replace_text_in_paragraph(paragraph)
    
    def insert_table_at_paragraph(self,paragraph, alignment=WD_TABLE_ALIGNMENT.CENTER):
        
        def set_table_border(table):
            tbl = table._element
            tbl_pr = tbl.tblPr

            # 创建边框元素
            tbl_borders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')  # 线条类型
                border.set(qn('w:sz'), '4')        # 线条大小
                border.set(qn('w:space'), '0')     # 间距
                border.set(qn('w:color'), '000000')  # 线条颜色
                tbl_borders.append(border)

            tbl_pr.append(tbl_borders)
        
        def set_column_width(table, column_idx, width):
            for row in table.rows:
                row.cells[column_idx].width = width
        def set_row_height(row, height):
            for cell in row.cells:
                cell.height = height

        def set_table_alignment(table, alignment):
            """
            设置表格的对齐方式
            alignment 参数可以是 WD_TABLE_ALIGNMENT.LEFT, WD_TABLE_ALIGNMENT.CENTER, WD_TABLE_ALIGNMENT.RIGHT
            """
            table.alignment = alignment
        
        table = self.doc.add_table(rows=1, cols=2)

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = self.tableDF.columns[0]
        hdr_cells[1].text = self.tableDF.columns[1]

        for i, row in self.tableDF.iterrows():
            row_cells = table.add_row().cells
            for j, cell_value in enumerate(row):
                if isinstance(cell_value, (int, float)):
                    row_cells[j].text = f"{cell_value:,.0f}"
                elif isinstance(cell_value, pd.Timestamp):
                    row_cells[j].text = cell_value.strftime('%d-%b-%Y')
                else:
                    row_cells[j].text = str(cell_value)

        set_table_border(table)
        set_table_alignment(table, alignment)

        # 设置列宽为4厘米
        set_column_width(table, 0, Cm(4))
        set_column_width(table, 1, Cm(4))

        # 设置行高为1厘米
        for row in table.rows:
            set_row_height(row, Cm(1))
        paragraph._element.addnext(table._element)

        def generateNewTermSheet(self):
            pass


class SDACGenerator(TSGenerator):
    def __init__(self,termPath = 'TermDetails\SDAC_template_field.xlsx'):
        self.termSheetPath = 'TemplateStore\SDAC_template.docx'
        self.termPath = termPath
        self.doc = Document(self.termSheetPath)
        self.fieldDF = pd.read_excel(self.termPath, sheet_name='Sheet1', usecols="A:B").dropna()
        self.tableDF = pd.read_excel(self.termPath, sheet_name='Sheet1', usecols="D:E").dropna(how='all')
        self.replaceDict  = self.getReplaceDict()

    def generateNewTermSheet(self,name=None):
        for paragraph in self.doc.paragraphs:
            self.replace_text_in_paragraph(paragraph)

        # 替换表格中的占位符，同时保留格式
        for table in self.doc.tables:
            self.replace_text_in_table(table)

        target_text = "{{ insert_table_here }}"
        for paragraph in self.doc.paragraphs:
            if target_text in paragraph.text:
                self.insert_table_at_paragraph(paragraph, alignment=WD_TABLE_ALIGNMENT.RIGHT)
                paragraph.text = ""
                break
        self.doc.save("TemplateResults\SDAC Template" + f" {name}.docx" if name else ".docx")